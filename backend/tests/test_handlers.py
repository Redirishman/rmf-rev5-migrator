"""End-to-end handler tests against moto-backed AWS."""

from __future__ import annotations

import io
import json

import pytest
from docx import Document as DocxDocument

from rmf_migrator.common.http import HttpError
from rmf_migrator.common.models import DocumentStatus, JobStatus
from rmf_migrator.handlers.create_project import _create
from rmf_migrator.handlers.enqueue_parse import _enqueue
from rmf_migrator.handlers.get_job import _get_job
from rmf_migrator.handlers.parse_document import run_parse_job
from rmf_migrator.handlers.request_upload import _request_upload


def _event(body=None, path=None, headers=None):
    return {
        "body": json.dumps(body) if body is not None else None,
        "pathParameters": path or {},
        "headers": headers or {},
    }


def _make_docx_bytes() -> bytes:
    doc = DocxDocument()
    doc.add_heading("Access Control Policy", level=1)
    doc.add_paragraph("The organization manages accounts.")
    doc.add_heading("AC-2 Account Management", level=2)
    doc.add_paragraph("Accounts are reviewed quarterly.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---- create_project --------------------------------------------------------


def test_create_project_happy(deps):
    resp = _create(_event(body={"name": "System Alpha"}), deps)
    assert resp["statusCode"] == 201
    project = json.loads(resp["body"])
    assert project["name"] == "System Alpha"
    assert project["baseline"] == "generic_800_53"
    assert deps.repo.get_project(project["project_id"]) is not None


def test_create_project_records_trusted_identity(deps):
    resp = _create(_event(body={"name": "Sys"}, headers={"X-Remote-User": "jdoe"}), deps)
    project = json.loads(resp["body"])
    assert project["created_by"] == "jdoe"


def test_create_project_anonymous_without_header(deps):
    resp = _create(_event(body={"name": "Sys"}), deps)
    assert json.loads(resp["body"])["created_by"] == "anonymous"


def test_create_project_requires_name(deps):
    with pytest.raises(HttpError) as exc:
        _create(_event(body={}), deps)
    assert exc.value.status == 400


def test_create_project_rejects_bad_baseline(deps):
    with pytest.raises(HttpError) as exc:
        _create(_event(body={"name": "S", "baseline": "nope"}), deps)
    assert exc.value.status == 400


# ---- request_upload --------------------------------------------------------


def test_request_upload_returns_presigned_url(deps):
    project = json.loads(_create(_event(body={"name": "S"}), deps)["body"])
    resp = _request_upload(
        _event(body={"filename": "ac-policy.docx"}, path={"project_id": project["project_id"]}),
        deps,
    )
    assert resp["statusCode"] == 201
    payload = json.loads(resp["body"])
    assert payload["upload"]["method"] == "PUT"
    assert "url" in payload["upload"]
    assert payload["upload"]["headers"]["x-amz-server-side-encryption"] == "aws:kms"
    assert payload["document"]["status"] == DocumentStatus.UPLOADED.value


def test_request_upload_rejects_non_docx(deps):
    project = json.loads(_create(_event(body={"name": "S"}), deps)["body"])
    with pytest.raises(HttpError) as exc:
        _request_upload(
            _event(body={"filename": "policy.pdf"}, path={"project_id": project["project_id"]}),
            deps,
        )
    assert exc.value.status == 400


def test_request_upload_404_missing_project(deps):
    with pytest.raises(HttpError) as exc:
        _request_upload(
            _event(body={"filename": "x.docx"}, path={"project_id": "proj_missing"}),
            deps,
        )
    assert exc.value.status == 404


def test_request_upload_increments_document_count(deps):
    project = json.loads(_create(_event(body={"name": "S"}), deps)["body"])
    pid = project["project_id"]
    for name in ("a.docx", "b.docx"):
        _request_upload(_event(body={"filename": name}, path={"project_id": pid}), deps)
    assert deps.repo.get_project(pid).document_count == 2


# ---- enqueue_parse ---------------------------------------------------------


def test_enqueue_parse_sends_sqs_and_creates_job(deps):
    project = json.loads(_create(_event(body={"name": "S"}), deps)["body"])
    pid = project["project_id"]
    up = json.loads(
        _request_upload(_event(body={"filename": "a.docx"}, path={"project_id": pid}), deps)["body"]
    )
    did = up["document"]["document_id"]

    resp = _enqueue(_event(path={"project_id": pid, "document_id": did}), deps)
    assert resp["statusCode"] == 202

    msgs = deps.sqs.receive_message(QueueUrl=deps.config.parse_queue_url).get("Messages", [])
    assert len(msgs) == 1
    body = json.loads(msgs[0]["Body"])
    assert body["document_id"] == did
    assert body["project_id"] == pid


def test_enqueue_parse_404_missing_document(deps):
    project = json.loads(_create(_event(body={"name": "S"}), deps)["body"])
    with pytest.raises(HttpError) as exc:
        _enqueue(
            _event(path={"project_id": project["project_id"], "document_id": "doc_missing"}),
            deps,
        )
    assert exc.value.status == 404


# ---- parse worker (end to end) ---------------------------------------------


def test_run_parse_job_parses_and_persists_sections(deps):
    project = json.loads(_create(_event(body={"name": "S"}), deps)["body"])
    pid = project["project_id"]
    up = json.loads(
        _request_upload(_event(body={"filename": "a.docx"}, path={"project_id": pid}), deps)["body"]
    )
    document = up["document"]
    did = document["document_id"]

    # Upload real .docx bytes to the (moto) S3 key the document points at.
    deps.store._s3.put_object(  # noqa: SLF001 — test reaching into the store's client
        Bucket=deps.config.documents_bucket,
        Key=document["s3_key"],
        Body=_make_docx_bytes(),
    )

    job = json.loads(_enqueue(_event(path={"project_id": pid, "document_id": did}), deps)["body"])[
        "job"
    ]

    run_parse_job(pid, did, job["job_id"], deps)

    updated_doc = deps.repo.get_document(pid, did)
    assert updated_doc.status == DocumentStatus.PARSED
    assert updated_doc.section_count == 2

    sections = deps.repo.list_sections(did)
    headings = {s.heading for s in sections}
    assert "Access Control Policy" in headings
    assert "AC-2 Account Management" in headings

    updated_job = deps.repo.get_job(pid, job["job_id"])
    assert updated_job.status == JobStatus.SUCCEEDED


def test_run_parse_job_marks_failed_on_bad_bytes(deps):
    project = json.loads(_create(_event(body={"name": "S"}), deps)["body"])
    pid = project["project_id"]
    up = json.loads(
        _request_upload(_event(body={"filename": "a.docx"}, path={"project_id": pid}), deps)["body"]
    )
    document = up["document"]
    did = document["document_id"]

    deps.store._s3.put_object(  # noqa: SLF001
        Bucket=deps.config.documents_bucket,
        Key=document["s3_key"],
        Body=b"not a real docx",
    )
    job = json.loads(_enqueue(_event(path={"project_id": pid, "document_id": did}), deps)["body"])[
        "job"
    ]

    with pytest.raises(Exception):  # noqa: B017 — re-raised for SQS retry semantics
        run_parse_job(pid, did, job["job_id"], deps)

    assert deps.repo.get_document(pid, did).status == DocumentStatus.FAILED
    failed_job = deps.repo.get_job(pid, job["job_id"])
    assert failed_job.status == JobStatus.FAILED
    assert failed_job.error_type is not None


# ---- get_job ---------------------------------------------------------------


def test_get_job_returns_status(deps):
    project = json.loads(_create(_event(body={"name": "S"}), deps)["body"])
    pid = project["project_id"]
    up = json.loads(
        _request_upload(_event(body={"filename": "a.docx"}, path={"project_id": pid}), deps)["body"]
    )
    did = up["document"]["document_id"]
    job = json.loads(_enqueue(_event(path={"project_id": pid, "document_id": did}), deps)["body"])[
        "job"
    ]

    resp = _get_job(_event(path={"project_id": pid, "job_id": job["job_id"]}), deps)
    assert resp["statusCode"] == 200
    assert json.loads(resp["body"])["job_id"] == job["job_id"]


def test_get_job_404(deps):
    project = json.loads(_create(_event(body={"name": "S"}), deps)["body"])
    with pytest.raises(HttpError) as exc:
        _get_job(_event(path={"project_id": project["project_id"], "job_id": "job_x"}), deps)
    assert exc.value.status == 404
