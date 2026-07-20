"""Export worker round-trip: draft text -> Rev 5 .docx in S3 (moto)."""

from __future__ import annotations

import io

import pytest
from docx import Document as DocxDocument

from rmf_migrator.common.models import (
    Document,
    DocumentStatus,
    Draft,
    DraftStatus,
    ExportJob,
    JobStatus,
    Project,
    Section,
)
from rmf_migrator.common.storage import build_export_key
from rmf_migrator.docx.parser import parse_docx_bytes
from rmf_migrator.handlers.worker import process_event, run_export_job
from tests.conftest import sqs_event


def _docx_bytes() -> bytes:
    doc = DocxDocument()
    doc.add_paragraph("Purpose statement.")  # preamble order 0
    doc.add_heading("Access Control Policy", level=1)  # order 1
    doc.add_paragraph("Old body.")
    doc.add_heading("AC-2 Account Management", level=2)  # order 2
    doc.add_paragraph("Old AC-2 body.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _seed(deps) -> tuple[str, str]:
    project = Project(name="Sys")
    deps.repo.put_project(project)
    pid = project.project_id
    document = Document(project_id=pid, filename="ac-policy.docx", s3_key="projects/x/doc.docx")
    document.status = DocumentStatus.EXPORTING
    deps.repo.put_document(document)
    did = document.document_id

    deps.store._s3.put_object(  # noqa: SLF001
        Bucket=deps.config.documents_bucket, Key=document.s3_key, Body=_docx_bytes()
    )
    sections = [
        Section(document_id=did, project_id=pid, order=1, level=1, heading="Access Control Policy"),
        Section(
            document_id=did, project_id=pid, order=2, level=2, heading="AC-2 Account Management"
        ),
    ]
    deps.repo.put_sections(sections)
    drafts = [
        Draft(
            project_id=pid,
            document_id=did,
            section_id=sections[0].section_id,
            order=1,
            draft_text="New Rev 5 AC policy text.",
            status=DraftStatus.APPROVED,
        ),
        Draft(
            project_id=pid,
            document_id=did,
            section_id=sections[1].section_id,
            order=2,
            draft_text="New AC-2 text.",
            edited_text="Edited AC-2 Rev 5 text.",
            status=DraftStatus.APPROVED,
        ),
    ]
    deps.repo.put_drafts(drafts)
    return pid, did


def test_run_export_job_writes_rev5_docx(deps):
    pid, did = _seed(deps)
    job = ExportJob(project_id=pid, document_id=did)
    deps.repo.put_export_job(job)

    run_export_job(pid, did, job.job_id, deps)

    document = deps.repo.get_document(pid, did)
    assert document.status == DocumentStatus.EXPORTED
    assert document.export_key == build_export_key(pid, did)
    assert deps.repo.get_export_job(pid, job.job_id).status == JobStatus.SUCCEEDED

    # Fetch the exported .docx from S3 and verify surgery applied edited/proposed text.
    out_bytes = deps.store.get_bytes(document.export_key)
    out = {s.order: s for s in parse_docx_bytes(out_bytes, document_id="d", project_id="p")}
    assert out[1].text == "New Rev 5 AC policy text."
    assert out[2].text == "Edited AC-2 Rev 5 text."  # edited wins over proposed
    assert "Purpose statement." in out[0].text  # preamble preserved


def test_run_export_job_keeps_original_text_for_empty_drafts(deps):
    # A draft whose effective text is empty (e.g. the failed-drafting fallback
    # sets draft_text="") must leave the original section untouched, not blank it.
    pid, did = _seed(deps)
    sections = deps.repo.list_sections(did)
    ac2 = next(s for s in sections if s.order == 2)
    draft = deps.repo.get_draft(did, ac2.section_id)
    draft.draft_text = ""
    draft.edited_text = None
    deps.repo.put_draft(draft)
    job = ExportJob(project_id=pid, document_id=did)
    deps.repo.put_export_job(job)

    run_export_job(pid, did, job.job_id, deps)

    document = deps.repo.get_document(pid, did)
    out_bytes = deps.store.get_bytes(document.export_key)
    out = {s.order: s for s in parse_docx_bytes(out_bytes, document_id="d", project_id="p")}
    assert out[2].text == "Old AC-2 body."  # original preserved, not deleted
    assert out[1].text == "New Rev 5 AC policy text."  # non-empty draft still applied


def test_export_dispatch_via_process_event(deps):
    pid, did = _seed(deps)
    job = ExportJob(project_id=pid, document_id=did)
    deps.repo.put_export_job(job)
    result = process_event(
        sqs_event({"kind": "export", "job_id": job.job_id, "project_id": pid, "document_id": did}),
        deps,
    )
    assert result["batchItemFailures"] == []
    assert deps.repo.get_document(pid, did).status == DocumentStatus.EXPORTED


def test_run_export_job_restores_status_on_failure(deps):
    pid, did = _seed(deps)
    # Point the document at a missing S3 key so export fails.
    document = deps.repo.get_document(pid, did)
    document.s3_key = "projects/x/missing.docx"
    deps.repo.put_document(document)
    job = ExportJob(project_id=pid, document_id=did)
    deps.repo.put_export_job(job)

    with pytest.raises(Exception):  # noqa: B017 — re-raised for SQS retry semantics
        run_export_job(pid, did, job.job_id, deps)

    # Document not left stuck in EXPORTING.
    assert deps.repo.get_document(pid, did).status == DocumentStatus.REVIEW_APPROVED
    assert deps.repo.get_export_job(pid, job.job_id).status == JobStatus.FAILED
