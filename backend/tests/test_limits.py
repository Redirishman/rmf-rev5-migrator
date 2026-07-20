"""Tests for the untrusted-.docx size guards (zip-bomb defense)."""

from __future__ import annotations

import io
import zipfile

import pytest
from docx import Document as DocxDocument

from rmf_migrator.common.limits import (
    MAX_COMPRESSION_RATIO,
    MAX_DOCX_BYTES,
    MAX_UNCOMPRESSED_BYTES,
    DocxTooLarge,
    ObjectTooLarge,
    guard_docx_bytes,
)
from rmf_migrator.docx.export_docx import export_rev5_docx
from rmf_migrator.docx.parser import parse_docx_bytes


def _real_docx() -> bytes:
    doc = DocxDocument()
    doc.add_heading("Access Control Policy", level=1)
    doc.add_paragraph("The organization manages accounts.")
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _zip_bomb() -> bytes:
    """A small zip whose single member decompresses enormously."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("word/document.xml", b" " * (80 * 1024 * 1024))
    return buf.getvalue()


def _lying_zip_bomb(actual_bytes: int = 80 * 1024 * 1024) -> bytes:
    """A zip bomb that under-declares its member's uncompressed size.

    Builds a real deflated member of ``actual_bytes``, then rewrites the
    uncompressed-size field in both the local file header (offset +22) and the
    central-directory header (offset +24) to a tiny value. A guard that trusts
    the declared sizes (``ZipInfo.file_size``) sees a small, low-ratio archive
    and lets it through.

    Note the exact failure mode downstream depends on the zip reader: CPython's
    ``zipfile`` caps member reads at the declared size and then fails the CRC
    check, while readers without that cap inflate the full ``actual_bytes``.
    The guard must reject the archive either way, without buffering it.
    """
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("word/document.xml", b" " * actual_bytes)
    raw = bytearray(buf.getvalue())

    forged = (4096).to_bytes(4, "little")

    lfh = raw.find(b"PK\x03\x04")
    if lfh != -1:
        raw[lfh + 22 : lfh + 26] = forged
    cdh = raw.find(b"PK\x01\x02")
    if cdh != -1:
        raw[cdh + 24 : cdh + 28] = forged

    return bytes(raw)


def test_real_document_passes_the_guard():
    guard_docx_bytes(_real_docx())  # must not raise


def test_zip_bomb_is_rejected_on_compression_ratio():
    data = _zip_bomb()
    assert len(data) < MAX_DOCX_BYTES  # small on the wire — that's the point
    with pytest.raises(DocxTooLarge):
        guard_docx_bytes(data)


def test_oversized_bytes_are_rejected():
    with pytest.raises(DocxTooLarge):
        guard_docx_bytes(b"x" * (MAX_DOCX_BYTES + 1))


def test_non_zip_bytes_are_rejected():
    with pytest.raises(DocxTooLarge):
        guard_docx_bytes(b"this is not a docx")


def test_ratio_constant_is_sane_for_real_documents():
    data = _real_docx()
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        total = sum(i.file_size for i in archive.infolist())
    assert total / len(data) < MAX_COMPRESSION_RATIO


def test_guard_rejects_a_zip_that_underdeclares_member_sizes():
    """The declared sizes in the zip's directory are attacker-controlled, so the
    guard must not trust them — it has to bound *actual* decompression."""
    data = _lying_zip_bomb()
    assert len(data) < MAX_DOCX_BYTES

    # Sanity: the forged archive looks tiny and low-ratio by its declared sizes,
    # which is exactly what a metadata-only guard would wave through.
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        declared_total = sum(info.file_size for info in archive.infolist())
    assert declared_total < MAX_UNCOMPRESSED_BYTES
    assert declared_total / len(data) < MAX_COMPRESSION_RATIO

    with pytest.raises(DocxTooLarge):
        guard_docx_bytes(data)


def test_parser_rejects_a_lying_zip_bomb():
    with pytest.raises(DocxTooLarge):
        parse_docx_bytes(_lying_zip_bomb(), document_id="d", project_id="p")


def test_parser_rejects_a_zip_bomb():
    with pytest.raises(DocxTooLarge):
        parse_docx_bytes(_zip_bomb(), document_id="d", project_id="p")


def test_export_rejects_a_zip_bomb():
    with pytest.raises(DocxTooLarge):
        export_rev5_docx(_zip_bomb(), {0: "replacement"})


def test_storage_download_enforces_the_streaming_byte_limit(deps):
    key = "projects/p/documents/d.docx"
    deps.store._s3.put_object(  # noqa: SLF001
        Bucket=deps.config.documents_bucket,
        Key=key,
        Body=b"1234",
    )

    with pytest.raises(ObjectTooLarge, match="exceeds"):
        deps.store.get_bytes(key, max_bytes=3)
