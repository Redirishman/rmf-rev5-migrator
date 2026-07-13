"""Round-trip tests for structure-preserving DOCX export (docx surgery).

Strategy: build a .docx, parse it to learn section orders, replace some section
bodies via the exporter, then re-parse the output and assert that headings/
structure are preserved and only the targeted bodies changed. The exporter's
section ordering must match the parser's, so unmapped sections stay put.
"""

from __future__ import annotations

import io

from docx import Document as DocxDocument

from rmf_migrator.docx.export_docx import export_rev5_docx
from rmf_migrator.docx.parser import parse_docx_bytes


def _build_docx() -> bytes:
    doc = DocxDocument()
    doc.add_paragraph("Purpose: this policy governs access.")  # preamble (order 0)
    doc.add_heading("Access Control Policy", level=1)  # order 1
    doc.add_paragraph("Old AC policy body.")
    doc.add_paragraph("More old AC text.")
    doc.add_heading("AC-2 Account Management", level=2)  # order 2
    doc.add_paragraph("Old AC-2 body.")
    doc.add_heading("AU-2 Audit Events", level=2)  # order 3
    doc.add_paragraph("Old AU-2 body.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _sections_by_order(data: bytes):
    return {s.order: s for s in parse_docx_bytes(data, document_id="d", project_id="p")}


def test_orders_match_parser():
    data = _build_docx()
    original = _sections_by_order(data)
    # Preamble + three headings.
    assert original[0].heading == ""
    assert original[1].heading == "Access Control Policy"
    assert original[2].heading == "AC-2 Account Management"
    assert original[3].heading == "AU-2 Audit Events"


def test_export_replaces_mapped_bodies_preserves_headings():
    data = _build_docx()
    new_bytes = export_rev5_docx(
        data,
        {
            1: "New Rev 5 access control policy language.",
            2: "New AC-2 language.\nAccounts reviewed monthly.",
        },
    )
    out = _sections_by_order(new_bytes)

    # Headings preserved, in order, with levels.
    assert out[1].heading == "Access Control Policy"
    assert out[1].level == 1
    assert out[2].heading == "AC-2 Account Management"
    assert out[3].heading == "AU-2 Audit Events"

    # Mapped bodies replaced.
    assert out[1].text == "New Rev 5 access control policy language."
    assert "New AC-2 language." in out[2].text
    assert "Accounts reviewed monthly." in out[2].text
    assert "Old AC-2 body." not in out[2].text
    assert "Old AC policy body." not in out[1].text


def test_export_leaves_unmapped_sections_unchanged():
    data = _build_docx()
    new_bytes = export_rev5_docx(data, {1: "Replaced AC only."})
    out = _sections_by_order(new_bytes)

    # AU-2 (order 3) not mapped -> original body preserved.
    assert out[3].text == "Old AU-2 body."
    # Preamble (order 0) not mapped -> preserved.
    assert "Purpose: this policy governs access." in out[0].text


def test_export_can_replace_preamble():
    data = _build_docx()
    new_bytes = export_rev5_docx(data, {0: "New preamble scope."})
    out = _sections_by_order(new_bytes)
    assert out[0].text == "New preamble scope."
    # Headings still intact after preamble surgery.
    assert out[1].heading == "Access Control Policy"


def test_export_multiparagraph_draft_splits_lines():
    data = _build_docx()
    new_bytes = export_rev5_docx(data, {2: "Line one.\nLine two.\nLine three."})
    # Re-open and count paragraphs under AC-2 before the next heading.
    doc = DocxDocument(io.BytesIO(new_bytes))
    texts = [p.text for p in doc.paragraphs]
    assert "Line one." in texts
    assert "Line two." in texts
    assert "Line three." in texts


def test_export_empty_draft_clears_body():
    data = _build_docx()
    new_bytes = export_rev5_docx(data, {3: ""})
    out = _sections_by_order(new_bytes)
    # AU-2 heading remains but its body is now empty.
    assert out[3].heading == "AU-2 Audit Events"
    assert out[3].text == ""


def test_export_output_is_valid_docx():
    data = _build_docx()
    new_bytes = export_rev5_docx(data, {1: "x"})
    # Should open without error and preserve the original preamble.
    doc = DocxDocument(io.BytesIO(new_bytes))
    assert any("Purpose:" in p.text for p in doc.paragraphs)
