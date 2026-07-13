"""Worker tests: mapping job + parse->map dispatch, against moto + fake Bedrock."""

from __future__ import annotations

import io
import json

from docx import Document as DocxDocument

from rmf_migrator.common.models import (
    Document,
    DocumentStatus,
    JobStatus,
    MappingJob,
    ParseJob,
    Section,
)
from rmf_migrator.handlers.worker import enqueue_mapping, process_event, run_mapping_job
from tests.conftest import FakeBedrock, sqs_event


def _docx_bytes() -> bytes:
    doc = DocxDocument()
    doc.add_heading("Access Control Policy", level=1)
    doc.add_paragraph("The organization manages accounts.")
    doc.add_heading("AC-2 Account Management", level=2)
    doc.add_paragraph("Accounts are reviewed quarterly.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _seed_parsed_document(deps) -> tuple[str, str]:
    """Create a project + document with two sections already persisted."""
    from rmf_migrator.common.models import Project

    project = Project(name="Sys")
    deps.repo.put_project(project)
    pid = project.project_id
    document = Document(project_id=pid, filename="a.docx", s3_key="k")
    document.status = DocumentStatus.PARSED
    deps.repo.put_document(document)
    did = document.document_id

    sections = [
        Section(document_id=did, project_id=pid, order=0, level=1, heading="AC Policy", text="x"),
        Section(document_id=did, project_id=pid, order=1, level=2, heading="AC-2", text="y"),
    ]
    deps.repo.put_sections(sections)
    return pid, did


def test_run_mapping_job_persists_mappings_and_status(deps):
    deps.bedrock = FakeBedrock()
    pid, did = _seed_parsed_document(deps)
    job = MappingJob(project_id=pid, document_id=did)
    deps.repo.put_mapping_job(job)

    run_mapping_job(pid, did, job.job_id, deps)

    assert deps.repo.get_document(pid, did).status == DocumentStatus.MAPPED
    done = deps.repo.get_mapping_job(pid, job.job_id)
    assert done.status == JobStatus.SUCCEEDED
    assert done.section_count == 2

    mappings = deps.repo.list_mappings(did)
    assert len(mappings) == 2
    assert mappings[0].proposed_control_ids == ["AC-2"]
    assert [m.order for m in mappings] == [0, 1]


def test_enqueue_mapping_creates_job_and_message(deps):
    pid, did = _seed_parsed_document(deps)
    job = enqueue_mapping(deps, pid, did)

    assert deps.repo.get_mapping_job(pid, job.job_id) is not None
    msgs = deps.sqs.receive_message(QueueUrl=deps.config.parse_queue_url).get("Messages", [])
    body = json.loads(msgs[0]["Body"])
    assert body["kind"] == "map"
    assert body["document_id"] == did


def test_parse_message_auto_chains_to_mapping(deps):
    """A parse SQS message parses the doc, then auto-enqueues a mapping job."""
    deps.bedrock = FakeBedrock()
    from rmf_migrator.common.models import Project

    project = Project(name="Sys")
    deps.repo.put_project(project)
    pid = project.project_id

    document = Document(project_id=pid, filename="a.docx", s3_key="projects/x/a.docx")
    deps.repo.put_document(document)
    did = document.document_id
    deps.store._s3.put_object(  # noqa: SLF001
        Bucket=deps.config.documents_bucket, Key=document.s3_key, Body=_docx_bytes()
    )
    parse_job = ParseJob(project_id=pid, document_id=did)
    deps.repo.put_job(parse_job)

    result = process_event(
        sqs_event(
            {"kind": "parse", "job_id": parse_job.job_id, "project_id": pid, "document_id": did}
        ),
        deps,
    )
    assert result["batchItemFailures"] == []

    # Parse completed...
    assert deps.repo.get_document(pid, did).status in (
        DocumentStatus.PARSED,
        DocumentStatus.MAPPING,
        DocumentStatus.MAPPED,
    )
    # ...and a mapping message was enqueued.
    msgs = deps.sqs.receive_message(
        QueueUrl=deps.config.parse_queue_url, MaxNumberOfMessages=10
    ).get("Messages", [])
    kinds = [json.loads(m["Body"])["kind"] for m in msgs]
    assert "map" in kinds


def test_map_message_dispatch_runs_mapping(deps):
    deps.bedrock = FakeBedrock()
    pid, did = _seed_parsed_document(deps)
    job = MappingJob(project_id=pid, document_id=did)
    deps.repo.put_mapping_job(job)

    process_event(
        sqs_event({"kind": "map", "job_id": job.job_id, "project_id": pid, "document_id": did}),
        deps,
    )
    assert deps.repo.get_document(pid, did).status == DocumentStatus.MAPPED
